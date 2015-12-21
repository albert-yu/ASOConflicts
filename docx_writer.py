import datetime
from docx import Document
from docx.shared import Inches
from parser import *

__author__ = 'Albert Yu'

# Date: 10-15-15


def create_document(date):
    """
    Creates the word document
    :param date, a datetime object
    :return: filename
    """
    document = Document()
    date_as_string = date.month.__str__() + "-" + date.day.__str__() + "-" + date.year.__str__()
    document.add_heading('CONFLICTS ' + date_as_string, 0)
    table = document.add_table(rows=1, cols=4)
    headings = table.rows[0].cells
    headings[0].text = 'Name'
    headings[1].text = 'Instrument'
    headings[2].text = 'Time'
    headings[3].text = 'Reason'
    file_name = 'Conflicts ' + date_as_string + '.docx'
    document.save(file_name)
    return file_name


def add_row(conflict, filename):
    """
    Add a row with all the conflict information to the table in the existing document
    :type conflict: Conflict object
    :param conflict: a conflict object from parser.py
    :param filename: string (contains ".docx" extension)
    :return: none
    """
    if (filename[len(filename) - 5: len(filename)] != '.docx') and \
       (filename[len(filename) - 4: len(filename)] != '.doc'):
        raise IOError
        print("File type must be a word document.")
    document = Document(filename)
    tables = document.tables()
    table = tables[0]
    row = table.add_row()
    row[0].text = conflict.name
    row[1].text = conflict.instrument
    row[2].text = conflict.time
    row[3].text = conflict.reason
    document.save(filename)  # this will overwrite the existing file, which is what we want


def main():
    doc_name = create_document(datetime.date.today())
    conflicts = store_conflicts('Conflicts.csv')
    for conflict in conflicts:
        add_row(conflict, doc_name)


if __name__ == "__main__":
    main()










