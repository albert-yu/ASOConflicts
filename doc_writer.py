from docx import Document
from parser import *

__author__ = 'Albert Yu'


def add_conflict_row(conflict, row):
    """
    Adds the conflict object to the given row
    :param conflict:
    :param row:
    :return: void
    """
    row[0].text = conflict.name
    row[1].text = conflict.instrument
    row[2].text = conflict.time
    row[3].text = conflict.reason


def create_headings(headings):
    """
    Creates the heading labels on the table
    :param headings: The cells of the first row of the table (e.g. table.rows[0].cells)
    :return: void
    """
    headings[0].text = 'Name'
    headings[1].text = 'Instrument'
    headings[2].text = 'Time'
    headings[3].text = 'Reason'


def create_doc(heading_pre, date_as_string):
    """
    Creates a document with an empty table and given heading
    :param heading_pre: string, the prefix of the heading (e.g. 'Conflicts ')
    :param date_as_string: string, the date as a string
    :return:
    """
    document = Document()
    document.add_heading(heading_pre + date_as_string, 0)
    table = document.add_table(rows=1, cols=4)
    headings = table.rows[0].cells
    create_headings(headings)
    return document


def main():
    date = datetime.date.today()
    date_as_string = date.month.__str__() + "-" + date.day.__str__() + "-" + date.year.__str__()
    document = create_doc('CONFLICTS ', date_as_string)
    file_name = 'Conflicts ' + date_as_string + '.docx'
    table = document.tables[0]

    # add ongoing conflicts first
    ongoing = store_conflicts('ongoing.csv')
    for conflict in ongoing:
        if conflict.date[:len(conflict.date) - 1] == int_to_weekday[date.weekday()]:
            row = table.add_row().cells
            add_conflict_row(conflict, row)

    # iterate through the list of conflicts and add them to the document table one by one
    conflicts = store_conflicts('Conflicts.csv')
    for conflict in conflicts:
        if string_to_datetime(conflict.date) == date:
            row = table.add_row().cells
            add_conflict_row(conflict, row)
    document.save(file_name)  # this will overwrite the existing file, which is what we want

if __name__ == "__main__":
    main()
